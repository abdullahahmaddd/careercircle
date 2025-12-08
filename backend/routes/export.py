from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from bson import ObjectId
from io import BytesIO
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.models import UserInDB
from backend.routes.auth import get_current_user
from backend.database import get_database

router = APIRouter()


def sanitize_filename(name: str) -> str:
    """Remove special characters from filename."""
    return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')


def create_ats_docx(resume_content: dict, role: str = "") -> BytesIO:
    """
    Create an ATS-compliant DOCX resume.
    Rules:
    - No tables, images, or text boxes
    - Standard fonts (Arial or Times New Roman)
    - Clear section headings
    - Consistent formatting
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Header with name
    name = resume_content.get('name', 'Candidate')
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_parts = []
    if resume_content.get('email'):
        contact_parts.append(resume_content['email'])
    if resume_content.get('phone'):
        contact_parts.append(resume_content['phone'])
    if resume_content.get('linkedin'):
        contact_parts.append(resume_content['linkedin'])
    
    if contact_parts:
        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    if resume_content.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(resume_content['summary'])
    
    # Experience
    experience = resume_content.get('experience', [])
    if experience:
        doc.add_heading('Work Experience', level=1)
        for exp in experience:
            title = exp.get('title', '')
            company = exp.get('company', '')
            start = exp.get('startDate', '')
            end = exp.get('endDate', '')
            
            # Job title and company
            job_heading = doc.add_paragraph()
            run_title = job_heading.add_run(f"{title}")
            run_title.bold = True
            if company:
                job_heading.add_run(f" at {company}")
            if start or end:
                job_heading.add_run(f" ({start} - {end})")
            
            # Bullet points
            description = exp.get('description', [])
            if isinstance(description, list):
                for bullet in description:
                    if bullet:
                        para = doc.add_paragraph(bullet, style='List Bullet')
            elif isinstance(description, str):
                doc.add_paragraph(description)
    
    # Education
    education = resume_content.get('education', [])
    if education:
        doc.add_heading('Education', level=1)
        for edu in education:
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            grad_date = edu.get('graduationDate', '')
            
            edu_para = doc.add_paragraph()
            run_degree = edu_para.add_run(degree)
            run_degree.bold = True
            if institution:
                edu_para.add_run(f" - {institution}")
            if grad_date:
                edu_para.add_run(f" (Graduated: {grad_date})")
    
    # Skills
    skills = resume_content.get('skills', [])
    if skills:
        doc.add_heading('Skills', level=1)
        skill_names = []
        for skill in skills:
            if isinstance(skill, dict):
                skill_names.append(skill.get('name', ''))
            elif isinstance(skill, str):
                skill_names.append(skill)
        
        skills_text = ', '.join([s for s in skill_names if s])
        doc.add_paragraph(skills_text)
    
    # Certifications
    certifications = resume_content.get('certifications', [])
    if certifications:
        doc.add_heading('Certifications', level=1)
        for cert in certifications:
            if isinstance(cert, dict):
                doc.add_paragraph(f"• {cert.get('name', '')} - {cert.get('issuer', '')}")
            elif isinstance(cert, str):
                doc.add_paragraph(f"• {cert}")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_ats_pdf(resume_content: dict, role: str = "") -> BytesIO:
    """
    Create an ATS-compliant PDF resume using reportlab.
    Text-based PDF with no embedded images.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation library not installed")
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
        textColor='black'
    )
    
    body_style = styles['Normal']
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        leftIndent=20,
        bulletIndent=10
    )
    
    story = []
    
    # Name
    name = resume_content.get('name', 'Candidate')
    story.append(Paragraph(name, title_style))
    
    # Contact
    contact_parts = []
    if resume_content.get('email'):
        contact_parts.append(resume_content['email'])
    if resume_content.get('phone'):
        contact_parts.append(resume_content['phone'])
    if resume_content.get('linkedin'):
        contact_parts.append(resume_content['linkedin'])
    
    if contact_parts:
        story.append(Paragraph(' | '.join(contact_parts), contact_style))
    
    story.append(Spacer(1, 12))
    
    # Summary
    if resume_content.get('summary'):
        story.append(Paragraph('PROFESSIONAL SUMMARY', section_style))
        story.append(Paragraph(resume_content['summary'], body_style))
    
    # Experience
    experience = resume_content.get('experience', [])
    if experience:
        story.append(Paragraph('WORK EXPERIENCE', section_style))
        for exp in experience:
            title = exp.get('title', '')
            company = exp.get('company', '')
            start = exp.get('startDate', '')
            end = exp.get('endDate', '')
            
            job_line = f"<b>{title}</b>"
            if company:
                job_line += f" at {company}"
            if start or end:
                job_line += f" ({start} - {end})"
            
            story.append(Paragraph(job_line, body_style))
            
            description = exp.get('description', [])
            if isinstance(description, list):
                for bullet in description:
                    if bullet:
                        story.append(Paragraph(f"• {bullet}", bullet_style))
            elif isinstance(description, str):
                story.append(Paragraph(description, body_style))
            
            story.append(Spacer(1, 6))
    
    # Education
    education = resume_content.get('education', [])
    if education:
        story.append(Paragraph('EDUCATION', section_style))
        for edu in education:
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            grad_date = edu.get('graduationDate', '')
            
            edu_line = f"<b>{degree}</b>"
            if institution:
                edu_line += f" - {institution}"
            if grad_date:
                edu_line += f" (Graduated: {grad_date})"
            
            story.append(Paragraph(edu_line, body_style))
    
    # Skills
    skills = resume_content.get('skills', [])
    if skills:
        story.append(Paragraph('SKILLS', section_style))
        skill_names = []
        for skill in skills:
            if isinstance(skill, dict):
                skill_names.append(skill.get('name', ''))
            elif isinstance(skill, str):
                skill_names.append(skill)
        
        skills_text = ', '.join([s for s in skill_names if s])
        story.append(Paragraph(skills_text, body_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


@router.get("/docx/{resume_id}")
async def export_resume_docx(
    resume_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    """Export a resume as an ATS-compliant DOCX file."""
    if not ObjectId.is_valid(resume_id):
        raise HTTPException(status_code=400, detail="Invalid resume ID")
    
    db = get_database()
    resume = await db.resumes.find_one({
        "_id": ObjectId(resume_id),
        "user_id": str(current_user.id)
    })
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    content = resume.get('content', {})
    name = sanitize_filename(content.get('name', 'Resume'))
    resume_name = sanitize_filename(resume.get('name', 'Resume'))
    
    filename = f"{name}_{resume_name}.docx"
    
    docx_buffer = create_ats_docx(content)
    
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/pdf/{resume_id}")
async def export_resume_pdf(
    resume_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    """Export a resume as an ATS-compliant PDF file."""
    if not ObjectId.is_valid(resume_id):
        raise HTTPException(status_code=400, detail="Invalid resume ID")
    
    db = get_database()
    resume = await db.resumes.find_one({
        "_id": ObjectId(resume_id),
        "user_id": str(current_user.id)
    })
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    content = resume.get('content', {})
    name = sanitize_filename(content.get('name', 'Resume'))
    resume_name = sanitize_filename(resume.get('name', 'Resume'))
    
    filename = f"{name}_{resume_name}.pdf"
    
    pdf_buffer = create_ats_pdf(content)
    
    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
