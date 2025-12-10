import io
import re
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from functools import lru_cache
from pypdf import PdfReader
from docx import Document

# ============================================================================
# PRE-COMPILED REGEX PATTERNS FOR EFFICIENCY
# ============================================================================

# Bullet point patterns
BULLET_PATTERNS = [
    re.compile(r'^[\s]*[•\-\*\u2022\u2023\u25E6\u2043\u2219○◦▪►→·▸▹‣⁃∙◘◙]\s+(.+)$'),
    re.compile(r'^[\s]*[\-\*\+]\s+(.+)$'),
    re.compile(r'^[\s]*\d+[\.\)]\s+(.+)$'),
    re.compile(r'^[\s]*[a-z][\.\)]\s+(.+)$'),
    re.compile(r'^[\s]*[ivxIVX]+[\.\)]\s+(.+)$'),
]

# Contact information patterns
EMAIL_PATTERN = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
PHONE_PATTERNS = [
    re.compile(r'\+?\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),
    re.compile(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
    re.compile(r'\d{10,15}'),
]
LINKEDIN_PATTERNS = [
    re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9-]+)', re.IGNORECASE),
    re.compile(r'linkedin\.com/([a-zA-Z0-9-]+)', re.IGNORECASE),
    re.compile(r'linkedin:?\s+([a-zA-Z0-9-]+)', re.IGNORECASE),
    re.compile(r'(?:^|\s)in/([a-zA-Z0-9-]+)(?:\s|$)'),
]

# Date patterns
DATE_PATTERN = re.compile(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4}|(?:19|20)\d{2}(?:\s*(?:-|–|—|to|TO)\s*(?:Present|Current|Now|(?:19|20)\d{2}))?|\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4}', re.IGNORECASE)
MONTH_YEAR_PATTERN = re.compile(r'([a-z]+)[\s,]+(\d{4})', re.IGNORECASE)
DAY_MONTH_YEAR_PATTERN = re.compile(r'(\d{1,2})[\s]+([a-z]+)[\s,]+(\d{4})', re.IGNORECASE)
DATE_SLASH_PATTERN = re.compile(r'(\d{1,2})[\.\/\-](\d{1,2})[\.\/\-](\d{2,4})')
YEAR_PATTERN = re.compile(r'\b(19|20)\d{2}\b')
QUARTER_PATTERN = re.compile(r'q[1-4]\s+(\d{4})', re.IGNORECASE)

# Multi-column detection
COLUMN_GAP_PATTERN = re.compile(r'\s{10,}')

# Skills patterns
BULLET_PREFIX_PATTERN = re.compile(r'^[\s]*[•\-\*\u2022○◦▪►→·]\s*')
SKILL_DELIMITER_PATTERN = re.compile(r'[,|•·;/]\s*')

# Common job titles to exclude from name detection
COMMON_JOB_TITLES = {
    'software engineer', 'senior software engineer', 'developer', 'programmer',
    'data scientist', 'analyst', 'manager', 'director', 'consultant',
    'engineer', 'architect', 'designer', 'lead', 'principal', 'staff',
    'intern', 'associate', 'specialist', 'coordinator', 'administrator'
}

# ============================================================================
# ENHANCED BULLET POINT DETECTION
# ============================================================================

def detect_bullet_points(text: str) -> List[str]:
    """
    Enhanced bullet point detection with pre-compiled regex patterns.
    Handles various bullet styles, numbering, and indentation.
    """
    lines = text.split('\n')
    bullets = []
    
    current_bullet = None
    
    for line in lines:
        if not line.strip():
            continue
            
        matched = False
        for pattern in BULLET_PATTERNS:
            match = pattern.match(line)
            if match:
                # Save previous bullet if exists
                if current_bullet:
                    bullets.append(current_bullet.strip())
                
                # Start new bullet
                current_bullet = match.group(1) if match.lastindex else match.group(0)
                matched = True
                break
        
        # If no bullet pattern matched but line is indented or starts lowercase, it's a continuation
        if not matched and current_bullet:
            stripped = line.strip()
            # Check if it's a continuation (indented OR starts with lowercase OR very short)
            if (re.match(r'^[\s]{2,}', line) or 
                (stripped and stripped[0].islower()) or
                (stripped and len(stripped) < 80 and not re.search(r'\d{4}', stripped))):
                current_bullet += ' ' + stripped
            else:
                # This is a new non-bullet line, save current and skip
                bullets.append(current_bullet.strip())
                current_bullet = None
    
    # Add last bullet
    if current_bullet:
        bullets.append(current_bullet.strip())
    
    return bullets


# ============================================================================
# COMPREHENSIVE DATE PARSING
# ============================================================================

class DateParser:
    """Enhanced date parser supporting multiple formats and languages."""
    
    MONTH_NAMES = {
        'jan': 1, 'january': 1, 'enero': 1,
        'feb': 2, 'february': 2, 'febrero': 2,
        'mar': 3, 'march': 3, 'marzo': 3,
        'apr': 4, 'april': 4, 'abril': 4,
        'may': 5, 'mayo': 5,
        'jun': 6, 'june': 6, 'junio': 6,
        'jul': 7, 'july': 7, 'julio': 7,
        'aug': 8, 'august': 8, 'agosto': 8,
        'sep': 9, 'sept': 9, 'september': 9, 'septiembre': 9,
        'oct': 10, 'october': 10, 'octubre': 10,
        'nov': 11, 'november': 11, 'noviembre': 11,
        'dec': 12, 'december': 12, 'diciembre': 12
    }
    
    MONTH_ABBREV = {v: k.capitalize()[:3] for k, v in {'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 
                    'may': 5, 'jun': 6, 'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}.items()}
    
    @classmethod
    @lru_cache(maxsize=128)
    def parse_date(cls, date_str: str) -> Optional[str]:
        """
        Parse date from various formats and return normalized format (Mon YYYY).
        Supported formats:
        - August 2024, Aug 2024, august 2024
        - 24.01.24, 24/1/25, 01-24-2024
        - 25 August 2025, 25 aug 2025
        - 2024, Q3 2024
        - Present, Current, Now
        """
        if not date_str:
            return None
            
        date_str = date_str.strip()
        date_lower = date_str.lower()
        
        # Handle special cases
        if any(word in date_lower for word in ['present', 'current', 'now', 'ongoing']):
            return 'Present'
        
        # Pattern 1: Month Name + Year (August 2024, Aug 2024)
        match = MONTH_YEAR_PATTERN.search(date_lower)
        if match:
            month_str, year = match.groups()
            if month_str in cls.MONTH_NAMES:
                month_num = cls.MONTH_NAMES[month_str]
                return f"{cls.MONTH_ABBREV[month_num]} {year}"
        
        # Pattern 2: Day Month Year (25 August 2025, 25 aug 2025)
        match = DAY_MONTH_YEAR_PATTERN.search(date_lower)
        if match:
            day, month_str, year = match.groups()
            if month_str in cls.MONTH_NAMES:
                month_num = cls.MONTH_NAMES[month_str]
                return f"{cls.MONTH_ABBREV[month_num]} {year}"
        
        # Pattern 3: dd.mm.yy or dd/mm/yy or dd-mm-yy
        match = DATE_SLASH_PATTERN.search(date_str)
        if match:
            day, month, year = match.groups()
            month = int(month)
            year = int(year)
            
            # Handle 2-digit years
            if year < 100:
                year = 2000 + year if year < 50 else 1900 + year
            
            if 1 <= month <= 12:
                return f"{cls.MONTH_ABBREV[month]} {year}"
        
        # Pattern 4: mm/dd/yyyy or mm-dd-yyyy (US format)
        match = re.search(r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})', date_str)
        if match:
            first, second, year = match.groups()
            first, second = int(first), int(second)
            
            # Determine if it's mm/dd or dd/mm
            if first > 12:  # Must be dd/mm
                month = second
            elif second > 12:  # Must be mm/dd
                month = first
            elif first <= 12 and second <= 12:  # Ambiguous - assume mm/dd (US format)
                month = first
            else:
                month = first
            
            if 1 <= month <= 12:
                return f"{cls.MONTH_ABBREV[month]} {int(year)}"
        
        # Pattern 5: Just year (2024)
        match = YEAR_PATTERN.search(date_str)
        if match:
            return match.group(0)
        
        # Pattern 6: Quarter notation (Q3 2024)
        match = QUARTER_PATTERN.search(date_lower)
        if match:
            return match.group(1)
        
        return date_str  # Return original if no pattern matched
    
    @classmethod
    def parse_date_range(cls, date_str: str) -> Tuple[str, str]:
        """
        Parse date range string into start and end dates.
        Handles: "Aug 2020 - Dec 2022", "2020-Present", "Aug 2020 to Dec 2022"
        """
        if not date_str:
            return ("", "Present")
        
        # Split on common separators
        separators = [' - ', ' – ', ' — ', ' to ', ' TO ', '–', '—', '-']
        parts = None
        
        for sep in separators:
            if sep in date_str:
                parts = date_str.split(sep, 1)
                break
        
        if not parts:
            # Single date - assume it's ongoing
            parsed = cls.parse_date(date_str)
            return (parsed or date_str, "Present")
        
        start_str = parts[0].strip()
        end_str = parts[1].strip() if len(parts) > 1 else "Present"
        
        start_date = cls.parse_date(start_str) or start_str
        end_date = cls.parse_date(end_str) or end_str
        
        return (start_date, end_date)


# ============================================================================
# MULTI-COLUMN RESUME DETECTION
# ============================================================================

def detect_columns(text: str) -> List[str]:
    """
    Detect and merge multi-column resume layouts.
    Uses heuristics to identify column boundaries.
    """
    lines = text.split('\n')
    
    # Detect if resume has columns by checking for:
    # 1. Lines with excessive spacing in the middle
    # 2. Multiple sections on the same line
    # 3. Inconsistent left margins
    
    column_indicators = 0
    processed_lines = []
    
    for line in lines:
        # Check for excessive mid-line spacing (likely columns)
        if COLUMN_GAP_PATTERN.search(line):
            column_indicators += 1
            # Split on large gaps
            parts = COLUMN_GAP_PATTERN.split(line)
            processed_lines.extend([p.strip() for p in parts if p.strip()])
        else:
            processed_lines.append(line)
    
    # If many column indicators found, we likely have a multi-column resume
    if column_indicators > 3:
        return processed_lines
    
    return lines


# ============================================================================
# ENHANCED PDF EXTRACTION WITH LAYOUT ANALYSIS
# ============================================================================

def extract_text_from_pdf(file_content: bytes) -> str:
    """Enhanced PDF extraction with better layout handling."""
    try:
        pdf = PdfReader(io.BytesIO(file_content))
        text = ""
        
        for page_num, page in enumerate(pdf.pages):
            extracted = page.extract_text()
            if extracted:
                # Detect and handle multi-column layouts
                lines = detect_columns(extracted)
                text += '\n'.join(lines) + "\n"
        
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Enhanced DOCX extraction."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += ' | '.join(row_text) + "\n"
        
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")


# ============================================================================
# ENHANCED CONTACT EXTRACTION
# ============================================================================

def extract_contact_info(text: str, lines: List[str]) -> Dict[str, str]:
    """Extract contact information with better patterns."""
    
    # Email pattern - more comprehensive
    email_match = EMAIL_PATTERN.search(text)
    email = email_match.group(0) if email_match else ""
    
    # Phone patterns - international support
    phone = ""
    for pattern in PHONE_PATTERNS:
        phone_match = pattern.search(text)
        if phone_match:
            candidate = phone_match.group(0)
            # Validate it's not a date or other number
            if not re.match(r'^\d{4}$', candidate) and len(candidate) >= 10:
                phone = candidate
                break
    
    # LinkedIn pattern
    linkedin = ""
    for pattern in LINKEDIN_PATTERNS:
        linkedin_match = pattern.search(text)
        if linkedin_match:
            linkedin = f"linkedin.com/in/{linkedin_match.group(1)}"
            break
    
    # Name extraction - improved heuristics
    name = "Unknown Candidate"
    for line in lines[:10]:  # Check first 10 lines
        line_clean = line.strip()
        
        # Skip if empty or too long
        if not line_clean or len(line_clean) > 60:
            continue
        
        # Skip lines with contact info
        if (EMAIL_PATTERN.search(line_clean) or
            re.search(r'\d{5,}', line_clean) or
            re.search(r'linkedin|github|http|www\.', line_clean, re.IGNORECASE) or
            re.search(r'^\+?\d[\d\s\-()]+$', line_clean)):
            continue
        
        # Skip common job titles and section headers
        if line_clean.lower() in COMMON_JOB_TITLES:
            continue
        
        # Must have at least 2 letters
        if not re.search(r'[A-Za-z]{2,}', line_clean):
            continue
        
        # Check if it looks like a name (2-4 words, capitalized)
        words = line_clean.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w):
                name = line_clean
                break
        elif len(words) == 1 and len(line_clean) > 3:
            # Single name (uncommon but possible)
            name = line_clean
            break
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin
    }


# ============================================================================
# ENHANCED SECTION DETECTION
# ============================================================================

def detect_section(line: str) -> Optional[str]:
    """Enhanced section header detection."""
    line_lower = line.lower().strip()
    line_clean = re.sub(r'[:\-–—|•=_]', '', line_lower).strip()
    
    # Skip if line is too long to be a header
    if len(line) > 80:
        return None
    
    section_keywords = {
        "summary": ["summary", "profile", "professional summary", "objective", 
                    "about me", "about", "career objective", "professional profile",
                    "career summary", "executive summary"],
        "experience": ["experience", "work experience", "work history", "employment", 
                       "professional experience", "career history", "employment history",
                       "work", "professional background", "relevant experience"],
        "education": ["education", "academic background", "qualifications", "academic",
                      "degrees", "academic qualifications", "educational background"],
        "skills": ["skills", "technical skills", "technologies", "core competencies",
                   "competencies", "expertise", "tools", "programming languages",
                   "technical proficiencies", "professional skills", "key skills"],
        "projects": ["projects", "personal projects", "key projects", "notable projects",
                     "academic projects", "professional projects"],
        "certifications": ["certifications", "certificates", "licenses", "credentials",
                           "professional certifications"],
    }
    
    # Check for exact or close matches
    for section, keywords in section_keywords.items():
        for keyword in keywords:
            if (keyword == line_clean or 
                line_clean.startswith(keyword) or 
                line_clean.endswith(keyword) or
                keyword in line_clean):
                return section
    
    return None


# ============================================================================
# ENHANCED EXPERIENCE PARSING
# ============================================================================

def parse_experience_section(lines: List[str]) -> List[Dict[str, Any]]:
    """Enhanced experience parsing with better date and bullet handling."""
    experiences = []
    
    if not lines:
        return experiences
    
    # Find all lines with dates
    date_indices = []
    for i, line in enumerate(lines):
        if DATE_PATTERN.search(line):
            date_indices.append(i)
    
    if not date_indices:
        # No dates found - try simpler parsing
        return parse_experience_fallback(lines)
    
    # Process each job entry
    for idx_pos, date_idx in enumerate(date_indices):
        # Get the block for this job (from date line to next date or end)
        end_idx = date_indices[idx_pos + 1] if idx_pos + 1 < len(date_indices) else len(lines)
        
        # Extract title, company, dates
        title = "Role"
        company = "Company"
        dates = ""
        
        # The line with the date usually contains the title too
        date_line = lines[date_idx].strip()
        date_match = DATE_PATTERN.search(date_line)
        if date_match:
            dates = date_match.group(0)
            # Extract title from the date line (text before the date)
            title_part = date_line[:date_match.start()].strip()
            if title_part and not any(pattern.match(title_part) for pattern in BULLET_PATTERNS):
                title = title_part
        
        # Parse dates
        start_date, end_date = DateParser.parse_date_range(dates)
        
        # Company is usually on the next non-bullet line after the date
        for i in range(date_idx + 1, end_idx):
            line_stripped = lines[i].strip()
            if (line_stripped and 
                not any(pattern.match(lines[i]) for pattern in BULLET_PATTERNS) and
                not DATE_PATTERN.search(line_stripped)):
                company = line_stripped
                break
        
        # Extract bullet points from the block
        block = lines[date_idx:end_idx]
        desc_text = '\n'.join(block)
        descriptions = detect_bullet_points(desc_text)
        
        # Filter out title, company, and date lines from descriptions
        descriptions = [d for d in descriptions 
                       if d not in [title, company] 
                       and not DATE_PATTERN.search(d)
                       and len(d) > 10]
        
        experiences.append({
            "title": title or "Role",
            "company": company or "Company",
            "startDate": start_date,
            "endDate": end_date,
            "description": descriptions[:15]  # Limit to 15 bullets
        })
    
    return experiences


def parse_experience_fallback(lines: List[str]) -> List[Dict[str, Any]]:
    """Fallback parser when no dates are found. Handles multiple formats."""
    experiences = []
    current_job = None
    i = 0
    
    while i < len(lines):
        line_stripped = lines[i].strip()
        
        # Skip empty or very short lines
        if not line_stripped or len(line_stripped) < 3:
            i += 1
            continue
        
        # Skip bullet points
        if any(pattern.match(lines[i]) for pattern in BULLET_PATTERNS):
            # If we have a current job, add this as a description
            if current_job:
                desc = detect_bullet_points(lines[i])[0] if detect_bullet_points(lines[i]) else None
                if desc and len(desc) > 10:
                    current_job["description"].append(desc)
            i += 1
            continue
        
        # Look for title/company patterns with separator
        if ' at ' in line_stripped or ' @ ' in line_stripped:
            if current_job:
                experiences.append(current_job)
            
            for sep in [' at ', ' @ ']:
                if sep in line_stripped:
                    parts = line_stripped.split(sep, 1)
                    current_job = {
                        "title": parts[0].strip(),
                        "company": parts[1].strip(),
                        "startDate": "",
                        "endDate": "Present",
                        "description": []
                    }
                    break
            i += 1
            continue
        
        # Look for title-company on consecutive lines (new logic)
        # If we find a non-bullet line followed by another non-bullet line,
        # treat first as title, second as company
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # Check if next line exists, is not a bullet, and is not too long
            if (next_line and 
                len(next_line) < 80 and
                not any(pattern.match(lines[i + 1]) for pattern in BULLET_PATTERNS)):
                
                # Save previous job if exists
                if current_job:
                    experiences.append(current_job)
                
                # Create new job
                current_job = {
                    "title": line_stripped,
                    "company": next_line,
                    "startDate": "",
                    "endDate": "Present",
                    "description": []
                }
                i += 2  # Skip both lines
                continue
        
        i += 1
    
    if current_job:
        experiences.append(current_job)
    
    return experiences


# ============================================================================
# ENHANCED EDUCATION PARSING
# ============================================================================

def parse_education_section(lines: List[str]) -> List[Dict[str, Any]]:
    """Enhanced education parsing."""
    education_entries = []
    current_entry = None
    
    degree_keywords = ["bachelor", "master", "phd", "doctor", "doctorate", "associate", 
                       "diploma", "b.s.", "m.s.", "b.a.", "m.a.", "mba", "b.tech", 
                       "m.tech", "bs", "ms", "ba", "ma", "bsc", "msc", "be", "me"]
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        line_lower = line_stripped.lower()
        has_degree = any(kw in line_lower for kw in degree_keywords)
        
        # Try to parse date
        parsed_date = DateParser.parse_date(line_stripped)
        has_date = parsed_date and parsed_date != line_stripped
        
        if has_degree or has_date:
            if current_entry and current_entry.get("degree"):
                education_entries.append(current_entry)
            
            degree = line_stripped
            institution = ""
            grad_date = ""
            
            # Extract date if present
            if has_date:
                grad_date = parsed_date
                degree = re.sub(r'\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', '', degree, flags=re.IGNORECASE).strip()
            
            # Try to separate degree and institution
            for sep in [' - ', ' at ', ' from ', ', ', ' | ', '\t']:
                if sep in line_stripped:
                    parts = line_stripped.split(sep, 1)
                    degree = parts[0].strip()
                    institution = parts[1].strip() if len(parts) > 1 else ""
                    break
            
            current_entry = {
                "degree": degree or "Degree",
                "institution": institution or "Institution",
                "graduationDate": grad_date or "N/A"
            }
        elif current_entry and not current_entry.get("institution"):
            # This line might be the institution
            current_entry["institution"] = line_stripped
    
    if current_entry and current_entry.get("degree"):
        education_entries.append(current_entry)
    
    return education_entries


# ============================================================================
# ENHANCED SKILLS PARSING
# ============================================================================

def parse_skills_section(lines: List[str]) -> List[Dict[str, str]]:
    """Enhanced skills parsing with better tokenization and categorization."""
    skills = []
    seen = set()
    
    text = ' '.join(line.strip() for line in lines)
    
    # Split by common delimiters but be smarter about it
    # First, try to detect if skills are on separate lines
    if any(BULLET_PREFIX_PATTERN.match(line) for line in lines):
        # Bullet-pointed skills
        potential_skills = []
        for line in lines:
            cleaned = BULLET_PREFIX_PATTERN.sub('', line).strip()
            if cleaned:
                potential_skills.append(cleaned)
    else:
        # Try splitting by em-dash first (common in paragraph-style skills)
        if '—' in text or '–' in text:
            # Split by em-dash or en-dash
            potential_skills = re.split(r'[—–]\s*', text)
        else:
            # Comma or delimiter-separated skills
            potential_skills = SKILL_DELIMITER_PATTERN.split(text)
    
    for skill in potential_skills:
        # Clean up
        clean_skill = skill.strip()
        clean_skill = re.sub(r'\([^)]*\)', '', clean_skill)  # Remove parenthetical
        clean_skill = BULLET_PREFIX_PATTERN.sub('', clean_skill)
        
        # Remove category prefixes like "Languages:", "Libraries:", "Tools:", etc.
        clean_skill = re.sub(r'^(?:\w+\s+)?(?:Languages?|Libraries?|Frameworks?|Tools?|Technologies?|Skills?):\s*', '', clean_skill, flags=re.IGNORECASE)
        
        # Split text further if it contains multiple "Name — description" patterns
        # E.g., "FastAPI —description, Automation — description" should become ["FastAPI", "Automation"]
        sub_skills = []
        if '—' in clean_skill or '–' in clean_skill:
            # This might have multiple skills separated by commas, each with em-dash
            # Split by comma first
            comma_parts = clean_skill.split(',')
            for part in comma_parts:
                part = part.strip()
                # Extract text before em-dash/en-dash
                if '—' in part or '–' in part:
                    skill_name = re.split(r'\s*[—–]\s*', part)[0].strip()
                    if skill_name:
                        sub_skills.append(skill_name)
                elif part and len(part) <= 50:  # Might be a skill without description
                    sub_skills.append(part)
        else:
            # No em-dash, treat as single skill
            sub_skills = [clean_skill]
        
        # Process each extracted skill
        for sk in sub_skills:
            sk = sk.strip()
            
            # Skip empty or invalid
            if not sk or len(sk) < 2:
                continue
            
            # Skip description fragments
            description_indicators = ['proficient', 'expertise', 'skilled', 'experience', 'knowledge', 'ability', 'familiar', 'using', 'enabling']
            if sk and sk[0].islower():
                continue
            if any(word in sk.lower() for word in description_indicators):
                continue
            
            # Normalize for deduplication
            skill_lower = re.sub(r'\s+\d+(\.\w*)?$', '', sk.lower())
            
            # Validate and add
            if (2 <= len(sk) <= 50 and 
                skill_lower not in seen and
                not re.match(r'^\d+$', sk) and
                not re.match(r'^[\W_]+$', sk)):
                skills.append({"name": sk, "level": ""})
                seen.add(skill_lower)
    
    return skills[:50]  # Limit to 50 skills


# ============================================================================
# MAIN PARSER
# ============================================================================

def parse_resume_text(text: str) -> Dict[str, Any]:
    """Main parsing function with enhanced logic."""
    
    # Handle multi-column layouts
    lines = detect_columns(text)
    raw_lines = [line for line in lines if line.strip()]
    
    # Extract contact info
    contact = extract_contact_info(text, raw_lines)
    
    # Section extraction
    sections = {
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }
    
    current_section = "summary"
    
    for line in raw_lines:
        detected = detect_section(line)
        if detected:
            current_section = detected
            continue
        sections[current_section].append(line)
    
    # Process sections
    experience = parse_experience_section(sections["experience"])
    education = parse_education_section(sections["education"])
    skills = parse_skills_section(sections["skills"])
    
    # Fallback for skills
    if not skills and sections["summary"]:
        skills = parse_skills_section(sections["summary"])
    
    summary = ' '.join(sections["summary"][:5])
    
    return {
        "name": contact["name"],
        "email": contact["email"],
        "phone": contact["phone"],
        "linkedin": contact["linkedin"],
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills
    }


async def parse_resume(file_content: bytes, filename: str) -> Dict[str, Any]:
    """Main entry point for resume parsing with enhanced error handling."""
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            text = extract_text_from_docx(file_content)
        else:
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = file_content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ValueError("Unsupported file format or encoding")
    except ValueError as e:
        raise e
    except Exception as e:
        raise ValueError(f"Failed to parse file: {str(e)}")
    
    if not text.strip():
        raise ValueError("Could not extract any text from the file")
    
    try:
        result = parse_resume_text(text)
        # Add validation - use filename as fallback for name
        if not result.get('name') or result['name'] == 'Unknown Candidate':
            candidate_name = filename.replace('.pdf', '').replace('.docx', '').replace('_', ' ').replace('-', ' ').strip()
            if candidate_name and len(candidate_name) < 60:
                result['name'] = candidate_name
        
        return result
    except Exception as e:
        # Return partial results with error indication
        return {
            "name": "Unknown Candidate",
            "email": "",
            "phone": "",
            "linkedin": "",
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "parsing_error": str(e)
        }