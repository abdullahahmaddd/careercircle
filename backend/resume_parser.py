import io
import re
from typing import Dict, Any, List, Optional
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")


def extract_contact_info(text: str, lines: List[str]) -> Dict[str, str]:
    """Extract name, email, phone, and LinkedIn from resume text."""
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    email = email_match.group(0) if email_match else ""
    
    # Phone patterns (various formats)
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{2,3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # International
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
        r'\d{10,11}',  # Plain 10-11 digits
    ]
    phone = ""
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0)
            break
    
    # LinkedIn pattern
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9-]+)'
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else ""
    
    # Name: First non-empty line that doesn't look like contact info
    name = "Unknown Candidate"
    for line in lines[:5]:  # Check first 5 lines
        line_clean = line.strip()
        # Skip if it looks like email, phone, address, or URL
        if (line_clean and 
            len(line_clean) < 60 and 
            not re.search(email_pattern, line_clean) and
            not re.search(r'\d{5,}', line_clean) and  # Zip codes
            not re.search(r'linkedin|github|http|www\.', line_clean, re.IGNORECASE) and
            not re.search(r'^\+?\d[\d\s\-()]+$', line_clean) and  # Phone
            re.search(r'[A-Za-z]{2,}', line_clean)):  # Has letters
            name = line_clean
            break
    
    return {"name": name, "email": email, "phone": phone, "linkedin": linkedin}


def detect_section(line: str) -> Optional[str]:
    """Detect if a line is a section header and return section type."""
    line_lower = line.lower().strip()
    line_clean = re.sub(r'[:\-–—|•]', '', line_lower).strip()
    
    section_keywords = {
        "summary": ["summary", "profile", "professional summary", "objective", "about me", "about", "career objective"],
        "experience": ["experience", "work experience", "work history", "employment", "professional experience", "career history", "employment history"],
        "education": ["education", "academic background", "qualifications", "academic", "degrees", "certifications"],
        "skills": ["skills", "technical skills", "technologies", "core competencies", "competencies", "expertise", "tools", "programming languages"],
        "projects": ["projects", "personal projects", "key projects"],
        "certifications": ["certifications", "certificates", "licenses"],
    }
    
    # Check if line is short enough to be a header (usually < 40 chars)
    if len(line) > 60:
        return None
    
    for section, keywords in section_keywords.items():
        for keyword in keywords:
            if keyword == line_clean or line_clean.startswith(keyword) or line_clean.endswith(keyword):
                return section
    
    return None


def parse_experience_section(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse experience section into structured job entries."""
    experiences: List[Dict[str, Any]] = []

    if not lines:
        return experiences

    # Patterns for job detection
    date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4}|(?:19|20)\d{2}(?:\s*(?:-|–|—|to)\s*(?:Present|Current|Now|(?:19|20)\d{2}))?|(?:\d{1,2}/\d{4})'
    title_company_pattern = r'^(.{2,80}?)\s*(?:[-–—|@,]\s*|\sat\s)(.{2,80})'

    # Heuristics helpers
    company_markers = ["inc", "llc", "ltd", "co", "corp", "corporation", "company", "gmbh", "plc", "s.a.", "sarl", "limited"]
    title_markers = ["engineer", "developer", "manager", "director", "lead", "senior", "sr\.", "jr\.", "consultant", "analyst", "officer", "architect", "designer", "vp", "vice"]

    def is_date_line(s: str) -> bool:
        return bool(re.search(date_pattern, s, re.IGNORECASE))

    def is_likely_company(s: str) -> bool:
        low = s.lower()
        if any(marker in low for marker in company_markers):
            return True
        # company names often contain '&' or 'and'
        if '&' in s or ' and ' in low:
            return True
        # all-caps short lines may be company
        if s.isupper() and 2 < len(s) < 60:
            return True
        return False

    def is_likely_title(s: str) -> bool:
        low = s.lower()
        if any(marker in low for marker in title_markers):
            return True
        # titles often contain role words and are sentence-cased
        words = s.split()
        if len(words) <= 6 and any(w[0].isupper() for w in words if w):
            return True
        return False

    # Helper to increment month in an end-date string (e.g., 'Feb 2027' -> 'Mar 2027')
    def increment_month_string(s: str) -> str:
        if not s:
            return s
        s = s.strip()
        # Handle Present/Current
        if re.search(r'Present|Current|Now', s, re.IGNORECASE):
            return s

        # Month name patterns
        month_map = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'sept': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }

        # Try MonthName YYYY
        m = re.search(r'([A-Za-z]+)\s+(19|20)\d{2}', s)
        if m:
            mon = m.group(1).lower()
            ym = re.search(r'(19|20)\d{2}', s)
            year = int(ym.group(0)) if ym else None
            if mon in month_map and year:
                month = month_map[mon]
                month += 1
                if month > 12:
                    month = 1
                    year += 1
                short_names = {1: 'Jan',2:'Feb',3:'Mar',4:'Apr',5:'May',6:'Jun',7:'Jul',8:'Aug',9:'Sep',10:'Oct',11:'Nov',12:'Dec'}
                return f"{short_names[month]} {year}"

        # Try mm/yyyy or m/yyyy
        m2 = re.search(r'(\d{1,2})[\/\-](19|20)\d{2}', s)
        if m2:
            mon = int(m2.group(1))
            year = int(re.search(r'(19|20)\d{2}', s).group(0))
            mon += 1
            if mon > 12:
                mon = 1
                year += 1
            return f"{mon:02d}/{year}"

        # If we only have a year, increment year (rare)
        y = re.search(r'^(19|20)\d{2}$', s.strip())
        if y:
            year = int(y.group(0)) + 1
            return str(year)

        # Otherwise, return original
        return s

    # Build blocks anchored by date lines. Include up to 2 lines above a date to capture title/company.
    indices = [i for i, ln in enumerate(lines) if is_date_line(ln)]

    if indices:
        for idx_pos, date_idx in enumerate(indices):
            start_idx = max(0, date_idx - 2)
            end_idx = indices[idx_pos + 1] if idx_pos + 1 < len(indices) else len(lines)
            block = [lines[i] for i in range(start_idx, end_idx)]

            # header candidates: first few non-empty lines in the block
            nonempty = [b.strip() for b in block if b.strip()]
            header_candidates = nonempty[:3]

            title = "Role"
            company = "Company"

            # Try to identify title/company
            if header_candidates:
                if len(header_candidates) >= 2:
                    a, b = header_candidates[0], header_candidates[1]
                    # prefer title/company ordering if heuristics match
                    if is_likely_title(a) and is_likely_company(b):
                        title, company = a, b
                    elif is_likely_company(a) and is_likely_title(b):
                        title, company = b, a
                    else:
                        # If company is on the next line (often italic below title),
                        # it's typically shorter and contains few words — treat that as company.
                        def short_company_candidate(s: str) -> bool:
                            ws = s.split()
                            if not s:
                                return False
                            if len(ws) <= 6 and len(s) < 60 and not re.search(r'\d{4}', s):
                                return True
                            return False

                        if short_company_candidate(b) and len(a.split()) > len(b.split()):
                            title, company = a, b
                        else:
                            # try regex split on ' at ' or separators
                            m = re.match(title_company_pattern, header_candidates[0], re.IGNORECASE)
                            if m:
                                title, company = m.group(1).strip(), m.group(2).strip()
                            else:
                                title, company = header_candidates[0], header_candidates[1]
                else:
                    # single header line
                    single = header_candidates[0]
                    m = re.match(title_company_pattern, single, re.IGNORECASE)
                    if m:
                        title, company = m.group(1).strip(), m.group(2).strip()
                    else:
                        # try to remove dates and use remaining as title
                        title = re.sub(date_pattern, '', single, flags=re.IGNORECASE).strip() or "Role"

            # Extract date string from block (first match)
            date_match = re.search(date_pattern, ' '.join(block), re.IGNORECASE)
            dates = date_match.group(0) if date_match else ""

            # Normalize date ranges
            start_date = ""
            end_date = ""
            if dates:
                parts = re.split(r'\s*(?:-|–|—|to)\s*', dates, maxsplit=1, flags=re.IGNORECASE)
                start_date = parts[0].strip() if parts else ""
                if len(parts) > 1:
                    raw_end = parts[1].strip()
                    end_date = increment_month_string(raw_end)
                else:
                    end_date = "Present" if re.search(r'Present|Current|Now', dates, re.IGNORECASE) else ""

            # Descriptions: take lines after header candidates in the block that look like bullets or sentences
            descriptions: List[str] = []
            # Determine how many header lines we consumed from the block
            consumed = 0
            # find first non-empty lines positions
            pos = 0
            for i, b in enumerate(block):
                if b.strip():
                    pos = i
                    break
            # Count header lines used
            for h in header_candidates:
                # try to find h in block sequentially
                for j in range(consumed, len(block)):
                    if block[j].strip() == h:
                        consumed = j + 1
                        break

            for l in block[consumed:]:
                raw_line = l
                if not raw_line.strip():
                    continue
                if is_date_line(raw_line):
                    continue
                bullet_start = bool(re.match(r'^[\s]*[•\-\*\u2022○◦▪►→·]', raw_line))
                indented = bool(re.match(r'^[\s]+\S', raw_line))
                clean_line = re.sub(r'^[\s]*[•\-\*\u2022○◦▪►→·]?\s*', '', raw_line).strip()
                if bullet_start:
                    if clean_line:
                        descriptions.append(clean_line)
                elif indented:
                    if clean_line:
                        if descriptions:
                            descriptions[-1] = descriptions[-1] + ' ' + clean_line
                        else:
                            descriptions.append(clean_line)
                else:
                    # Non-bullet line: treat as a description sentence only if reasonably long
                    if clean_line and len(clean_line) > 20:
                        descriptions.append(clean_line)

            experiences.append({
                "title": title or "Role",
                "company": company or "Company",
                "startDate": start_date,
                "endDate": end_date or "Present",
                "description": descriptions
            })
    else:
        # fallback: earlier line-based heuristic when no date anchors found
        current_job: Dict[str, Any] = {}
        current_descriptions: List[str] = []
        title_company_pattern_relaxed = title_company_pattern
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            title_match = re.match(title_company_pattern_relaxed, line_stripped, re.IGNORECASE)
            if title_match:
                # If there's an active job, finalize it before starting a new one
                if current_job:
                    current_job["description"] = current_descriptions
                    experiences.append(current_job)
                # start a new job
                title = title_match.group(1).strip()
                company = title_match.group(2).strip()
                current_job = {"title": title, "company": company, "startDate": "", "endDate": "Present"}
                current_descriptions = []
                continue
            if current_job:
                # same bullet logic as before
                raw_line = line
                bullet_start = bool(re.match(r'^[\s]*[•\-\*\u2022○◦▪►→·]', raw_line))
                indented = bool(re.match(r'^[\s]+\S', raw_line))
                clean_line = re.sub(r'^[\s]*[•\-\*\u2022○◦▪►→·]?\s*', '', raw_line).strip()
                if bullet_start:
                    if clean_line:
                        current_descriptions.append(clean_line)
                elif indented:
                    if clean_line:
                        if current_descriptions:
                            current_descriptions[-1] = current_descriptions[-1] + ' ' + clean_line
                        else:
                            current_descriptions.append(clean_line)
                else:
                    if clean_line and len(clean_line) > 20:
                        current_descriptions.append(clean_line)
        if current_job:
            current_job["description"] = current_descriptions
            experiences.append(current_job)

    return experiences


def parse_education_section(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse education section into structured entries."""
    education_entries = []
    current_entry: Dict[str, Any] = {}
    
    degree_keywords = ["bachelor", "master", "phd", "doctor", "associate", "diploma", "b.s.", "m.s.", "b.a.", "m.a.", "mba", "b.tech", "m.tech", "bs", "ms", "ba", "ma"]
    year_pattern = r'(?:19|20)\d{2}'
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        line_lower = line_stripped.lower()
        has_degree = any(kw in line_lower for kw in degree_keywords)
        has_year = bool(re.search(year_pattern, line_stripped))
        
        if has_degree or has_year:
            if current_entry and current_entry.get("degree"):
                education_entries.append(current_entry)
            
            # Extract year
            year_match = re.search(year_pattern, line_stripped)
            year = year_match.group(0) if year_match else ""
            
            # Try to separate degree and institution
            degree = line_stripped
            institution = ""
            
            # Common separators
            for sep in [" - ", " at ", " from ", ", ", " | "]:
                if sep in line_stripped:
                    parts = line_stripped.split(sep)
                    degree = parts[0].strip()
                    institution = parts[1].strip() if len(parts) > 1 else ""
                    break
            
            current_entry = {
                "degree": re.sub(year_pattern, '', degree).strip(),
                "institution": institution or "Institution",
                "graduationDate": year or "N/A"
            }
        elif current_entry:
            # Additional info for current entry
            if not current_entry.get("institution") or current_entry.get("institution") == "Institution":
                current_entry["institution"] = line_stripped
    
    if current_entry and current_entry.get("degree"):
        education_entries.append(current_entry)
    
    # Fallback if nothing parsed
    if not education_entries and lines:
        combined = " ".join(line.strip() for line in lines[:3] if line.strip())
        education_entries.append({
            "degree": combined[:100] if combined else "Education details in resume",
            "institution": "",
            "graduationDate": ""
        })
    
    return education_entries


def parse_skills_section(lines: List[str]) -> List[Dict[str, str]]:
    """Parse skills section into list of skills."""
    skills = []
    seen = set()
    
    # Combine all lines
    text = " ".join(line.strip() for line in lines)
    
    # Split by common delimiters
    potential_skills = re.split(r'[,|•·:\n/]', text)
    
    for skill in potential_skills:
        # Clean up the skill
        clean_skill = re.sub(r'^\s*[-–—*○◦▪►]\s*', '', skill)
        clean_skill = re.sub(r'\([^)]*\)', '', clean_skill)  # Remove parenthetical
        clean_skill = clean_skill.strip()
        
        # Validate skill
        if (clean_skill and 
            2 < len(clean_skill) < 50 and 
            clean_skill.lower() not in seen and
            not re.match(r'^\d+$', clean_skill)):  # Not just numbers
            skills.append({"name": clean_skill, "level": ""})
            seen.add(clean_skill.lower())
    
    return skills[:30]  # Limit to 30 skills


def parse_resume_text(text: str) -> Dict[str, Any]:
    """Parse resume text into structured data."""
    # Preserve raw lines (to keep indentation and bullet markers) and a stripped version for header/contact heuristics
    raw_lines = [line for line in text.split('\n') if line.strip()]
    stripped_lines = [line.strip() for line in raw_lines]

    # Extract contact info using stripped lines
    contact = extract_contact_info(text, stripped_lines)

    # Section extraction using state machine; store raw lines per section to preserve bullets/indent
    sections = {
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }

    current_section = "summary"

    # Use both raw and stripped lines to detect headers but keep raw for content
    for raw, stripped in zip(raw_lines, stripped_lines):
        detected = detect_section(stripped)
        if detected:
            current_section = detected
            continue
        sections[current_section].append(raw)
    
    # Process each section
    # Pass raw experience lines (with indentation) so bullets and continuations can be detected
    experience = parse_experience_section(sections["experience"])
    education = parse_education_section(sections["education"])
    skills = parse_skills_section(sections["skills"])
    
    # If skills are empty, try to extract from other sections
    if not skills and sections["summary"]:
        skills = parse_skills_section(sections["summary"][:3])
    
    summary = " ".join(sections["summary"][:5])  # First 5 lines of summary
    
    return {
        "name": contact["name"],
        "email": contact["email"],
        "phone": contact["phone"],
        "linkedin": contact["linkedin"],
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills
    }


async def parse_resume(file_content: bytes, filename: str) -> Dict[str, Any]:
    """Main entry point for resume parsing."""
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            text = extract_text_from_docx(file_content)
        else:
            # Try decoding as text
            try:
                text = file_content.decode('utf-8')
            except Exception:
                raise ValueError("Unsupported file format and could not decode as text")
    except ValueError as e:
        raise e
    except Exception as e:
        raise ValueError(f"Failed to parse file: {str(e)}")
            
    if not text.strip():
        raise ValueError("Could not extract any text from the file")

    return parse_resume_text(text)